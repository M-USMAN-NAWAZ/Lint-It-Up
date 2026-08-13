from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Lint-It-Up-User-Guidance.docx"


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Aptos Display"
    return heading


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()
    return table


def build_document():
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lint It Up\nUser Guidance")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Aptos Display"

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Experience flow, gameplay, controls, and user journey").italic = True

    add_heading(document, "Purpose", 1)
    add_body(
        document,
        "Lint It Up is a VR football training experience. The user begins in a theater-style playbook environment, "
        "reviews video and coaching panels, answers setup questions, enters an on-field scenario, completes a guided "
        "football play, and returns to the theater to review their saved answers."
    )

    add_heading(document, "Scenes And Main Flow", 1)
    add_table(
        document,
        ["Scene", "Role in the experience", "Main scripts"],
        [
            [
                "Theater",
                "Main hub for intro video, playbook panels, coach videos, voice Q&A, and post-game answer display.",
                "VideoPlayerTimelineControls, VideoPanelSwitcher, MetaVoiceTMPButton, TheaterGame1AnswerDisplay, SceneChange",
            ],
            [
                "Game 1",
                "Pre-play analysis scene. The user sees the formation, answers two football-read questions, then moves into the field gameplay scene.",
                "Game1QuestionnaireUI, ScenarioHutHutTimer, SceneOrbitCamera",
            ],
            [
                "Game",
                "Interactive VR football scenario. The user catches, carries, fakes, moves, and throws the ball to complete the play.",
                "VRFootballScenarioController, VRFootballScenarioUI, FormationRunController, ScenarioHutHutTimer",
            ],
        ],
    )

    add_heading(document, "High-Level User Journey", 1)
    add_numbered(
        document,
        [
            "The user starts in the Theater scene.",
            "The intro video panel can play automatically or be opened from the UI.",
            "The user navigates playbook and coach-video panels using world-space UI buttons.",
            "The user starts the football scenario flow from the Theater, which loads Game 1.",
            "Game 1 displays a play clock, pauses the camera at key moments, and asks the user two read-and-react questions.",
            "After the questions, the project loads the Game scene.",
            "The Game scene starts the VR football scenario with a countdown and on-field task prompts.",
            "If the user completes every task, the success flow returns to Theater and opens the scorecard/answer panel.",
            "If the user fails or times out, the relevant scene restarts or the timeout UI appears.",
        ],
    )

    add_heading(document, "Theater Hub Flow", 1)
    add_body(
        document,
        "The Theater scene is the front-end hub. It contains multiple UI panels under the main Content object. "
        "Most navigation uses Unity Button OnClick events that call GameObject.SetActive(true) or SetActive(false) "
        "to show and hide panels."
    )
    add_bullets(
        document,
        [
            "IntroVideoPanel is active in the scene and has a VideoPanelSwitcher.",
            "When the intro video finishes, VideoPanelSwitcher turns IntroVideoPanel off and turns Panel 1 on.",
            "Video panels such as Panel 7, Panel 8, Panel 10, Panel 11, Panel 12, Panel 13, and Panel 14 include VideoPlayerTimelineControls.",
            "VideoPlayerTimelineControls handles preparation, play/pause, timeline slider seeking, time display, and auto-hiding controls.",
            "Most video panels are opened by buttons rather than by an automatic video-completion chain.",
        ],
    )

    add_heading(document, "Video Panel Behavior", 2)
    add_numbered(
        document,
        [
            "A UI button is clicked with the XR ray/interactor.",
            "The button's OnClick event activates a target video panel.",
            "When the panel becomes active, VideoPlayerTimelineControls subscribes to the play/pause button and slider.",
            "The video is prepared if prepareVideoOnStart is enabled.",
            "The play/pause button toggles the VideoPlayer.",
            "The slider can seek through the video by setting normalized video time.",
            "Controls fade out after the configured visible time and reappear when used.",
        ],
    )

    add_heading(document, "Voice Question Flow", 1)
    add_body(
        document,
        "The Theater includes a voice interaction powered by Meta Voice SDK/Wit. The scene object named "
        "Voicerechognization provides the voice service, while the Content object stores the expected question and response "
        "through the MetaVoiceTMPButton component."
    )
    add_bullets(
        document,
        [
            "The Speak button activates listening through VoiceService.",
            "Partial and full transcriptions are received from the voice system.",
            "The transcription is normalized by lowercasing it and removing punctuation.",
            "If the user asks the configured New England / 2 Tight Ends Trips / adjustment-counter question, the configured answer appears.",
            "If the speech is not recognized as the expected question, the UI shows: This AI feature is coming soon.",
        ],
    )

    add_heading(document, "Game 1 Analysis Flow", 1)
    add_body(
        document,
        "Game 1 is a pre-play recognition and questionnaire sequence. It is designed to make the user observe the defensive look "
        "before entering the active football play."
    )
    add_numbered(
        document,
        [
            "Game1QuestionnaireUI starts a 40-second ScenarioHutHutTimer.",
            "The play clock appears in world space.",
            "After an initial delay, the orbit camera pauses and Question 1 appears.",
            "Question 1 asks: What do you see?",
            "The user chooses Option A, B, or C. The selected text is saved in PlayerPrefs.",
            "The UI hides, the orbit camera resumes, and player/highlighter visuals can appear.",
            "After the second delay, the orbit camera pauses again and Question 2 appears.",
            "Question 2 asks: What might they be giving us?",
            "The final answer is saved, the questionnaire hides, and the next scene loads.",
        ],
    )

    add_heading(document, "Question Defaults", 2)
    add_table(
        document,
        ["Question", "Prompt", "Options"],
        [
            [
                "1",
                "What do you see?",
                "A: SS in the box - Showing Cover 1; B: 5 Down & Possible Blitzer; C: Ben to provide",
            ],
            [
                "2",
                "What might they be giving us?",
                "A: Run; B: TE Release; C: Ben to provide",
            ],
        ],
    )

    add_heading(document, "On-Field Gameplay Flow", 1)
    add_body(
        document,
        "The Game scene uses VRFootballScenarioController as the main gameplay state machine. It prepares the formation, "
        "waits for the start condition, shows a countdown, launches the play, and checks each task in order."
    )
    add_table(
        document,
        ["Step", "Gameplay task", "What the user does", "Completion check"],
        [
            [
                "1",
                "Catch The Ball",
                "Catch the pass from the passer after the hut-hut timing.",
                "The football becomes held by the user.",
            ],
            [
                "2",
                "Run To Faker",
                "Carry the ball toward the highlighted faker spot.",
                "The user's body reaches the target radius while holding the ball.",
            ],
            [
                "3",
                "Fake The Hand Off",
                "Stretch the hand holding the ball toward the fake target.",
                "Either hand reaches the hand-target radius while the ball is held.",
            ],
            [
                "4",
                "Run To Safety",
                "Move to the safe spot before the defense closes.",
                "The user reaches the safe zone while holding the ball.",
            ],
            [
                "5",
                "Throw To Goal",
                "Release the ball toward the goal player with a throwing motion.",
                "The goal throw completes, the receiver catch triggers, or the ball reaches the target.",
            ],
        ],
    )

    add_heading(document, "Gameplay Rules And Failure", 1)
    add_bullets(
        document,
        [
            "The scenario begins with a countdown, defaulting to 3, 2, 1.",
            "The football is locked to its pass origin until the catch sequence starts.",
            "Objective indicators and hand indicators point the user to the next required target.",
            "Some tasks require the ball to remain in the user's hand.",
            "If a task timer expires or an early failure condition happens, the failure panel appears and the active scene reloads.",
            "If the user completes all tasks, the success panel displays TOUCHDOWN, then the project loads Theater.",
            "Before returning to Theater, the game sets a flag so Theater opens the scorecard/answer display automatically.",
        ],
    )

    add_heading(document, "Running Back Choice", 1)
    add_body(
        document,
        "During the running-back fake/pass segment, the project can present a tactical choice: keep holding the triggers to fake the handoff, "
        "or release near the running back to pass the ball to him. If the running-back pass branch is triggered, the running back can be chased "
        "and tackled, which leads into the failure/restart flow."
    )

    add_heading(document, "Timeouts", 1)
    add_bullets(
        document,
        [
            "Game 1 starts a 40-second timer.",
            "If Game 1 times out, Game 1 reloads.",
            "The Game scene can start a shorter retry timer after the football play begins or after failure handling.",
            "If the Game scene times out, the Game Timeout UI appears, optional audio plays, and the Game scene reloads.",
            "The timer stops when hut-hut/play action has begun through ScenarioHutHutTimer.StopForHutHut().",
        ],
    )

    add_heading(document, "Returning To Theater", 1)
    add_bullets(
        document,
        [
            "The right controller secondary button can exit back to Theater through RightControllerSceneExit.",
            "After a successful football scenario, TheaterGame1AnswerDisplay detects the return flag.",
            "Theater hides IntroVideoPanel, opens Panel 2, activates the scorecard hierarchy, and loads saved Game 1 answers.",
            "The saved answers come from PlayerPrefs keys Game1.Question1.AnswerText and Game1.Question2.AnswerText.",
        ],
    )

    add_heading(document, "User Controls Summary", 1)
    add_table(
        document,
        ["Action", "How the user performs it"],
        [
            ["Navigate Theater UI", "Aim at UI buttons with the XR interactor/ray and select."],
            ["Play or pause a video", "Open a video panel, then press its play/pause button."],
            ["Scrub a video", "Use the timeline slider on the active video panel."],
            ["Use voice Q&A", "Press Speak, ask the configured question, then read the response text."],
            ["Answer Game 1 questions", "Select Option A, B, or C on each questionnaire panel."],
            ["Catch/hold ball", "Grab or pinch the football with a VR hand/controller."],
            ["Complete movement tasks", "Move the body or hand to the highlighted target while obeying ball-holding requirements."],
            ["Throw to goal", "Release the ball toward the goal target with enough speed and direction."],
            ["Exit to Theater", "Use the right controller secondary button if the exit handler is active."],
        ],
    )

    add_heading(document, "Troubleshooting Notes", 1)
    add_bullets(
        document,
        [
            "If a video panel opens but does not play, check that its VideoPlayer reference is assigned and that the VideoClip exists.",
            "If video controls do not respond, check the playPauseButton, timelineSlider, and CanvasGroup references on VideoPlayerTimelineControls.",
            "If the intro video does not advance, verify VideoPanelSwitcher has panelToTurnOff set to IntroVideoPanel and panelToTurnOn set to Panel 1.",
            "If voice recognition says Voice service missing, verify the Meta Voice/Wit object in Theater is active and assigned.",
            "If the user returns to Theater without seeing answers, check that the win scene is Theater and the PlayerPrefs return flag is being set.",
            "If gameplay restarts unexpectedly, inspect task durations, completion radii, ball-in-hand requirements, and ScenarioHutHutTimer settings.",
        ],
    )

    add_heading(document, "Developer Reference", 1)
    add_table(
        document,
        ["Area", "Primary file"],
        [
            ["Theater video controls", "Assets/Scripts/VideoPlayerTimelineControls.cs"],
            ["Video completion panel switching", "Assets/Scripts/VideoPanelSwitcher.cs"],
            ["Theater voice Q&A", "Assets/Scripts/MetaVoiceTMPButton.cs"],
            ["Pre-play questionnaire", "Assets/Scripts/Game1QuestionnaireUI.cs"],
            ["On-field gameplay controller", "Assets/Scripts/VRFootballScenarioController.cs"],
            ["Gameplay UI panels", "Assets/Scripts/VRFootballScenarioUI.cs"],
            ["Formation animation and runners", "Assets/Scripts/FormationRunController.cs"],
            ["Timeout display/retry flow", "Assets/Scripts/ScenarioHutHutTimer.cs"],
            ["Post-game answer display", "Assets/Scripts/TheaterGame1AnswerDisplay.cs"],
            ["Scene exit shortcut", "Assets/Scripts/RightControllerSceneExit.cs"],
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
