from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Usman\Line-it-up")
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "Local-Unity-Android-Build-Plan.docx"


def set_page_margins(section, margin=0.75):
    inches = Inches(margin)
    section.top_margin = inches
    section.bottom_margin = inches
    section.left_margin = inches
    section.right_margin = inches


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25 * level)
    paragraph.add_run(text)
    return paragraph


def add_code_block(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shade_paragraph(p, "F4F7FA")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    return heading


doc = Document()
section = doc.sections[0]
set_page_margins(section)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Local Unity Android Build Automation Plan")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(24, 54, 93)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = subtitle.add_run("Local-only, Android-only, saved on your machine")
sub.italic = True
sub.font.size = Pt(10.5)
sub.font.color.rgb = RGBColor(90, 90, 90)

callout = doc.add_paragraph()
callout.paragraph_format.space_before = Pt(10)
callout.paragraph_format.space_after = Pt(10)
shade_paragraph(callout, "EAF2FF")
r = callout.add_run(
    "Important: this plan does not use GitHub-hosted build resources. "
    "Unity still performs the build locally in batch mode."
)
r.bold = True

add_heading(doc, "Requirements", 1)
for item in [
    "Local-only build process",
    "Android only",
    "APK saved locally on the system",
    "Button-triggered first, with optional git automation later",
]:
    add_bullet(doc, item)

add_heading(doc, "Recommended Design", 1)
for item in [
    "A Unity build script in Assets/Editor/BuildScript.cs",
    "A local PowerShell build script in BuildTools/Build-Android.ps1",
    "A one-click button launcher in BuildTools/Build-Android.bat",
    "An optional local git hook if automatic local enforcement is needed",
]:
    add_bullet(doc, item)

add_heading(doc, "Recommended Folder Structure", 1)
add_code_block(
    doc,
    """
Assets/
  Editor/
    BuildScript.cs

BuildTools/
  Build-Android.ps1
  Build-Android.bat

Builds/
  Android/

Logs/
""",
)

add_heading(doc, "Unity Build Script Template", 1)
add_code_block(
    doc,
    """
using System;
using System.Linq;
using UnityEditor;
using UnityEditor.Build.Reporting;

public static class BuildScript
{
    private static string[] GetEnabledScenes()
    {
        return EditorBuildSettings.scenes
            .Where(scene => scene.enabled)
            .Select(scene => scene.path)
            .ToArray();
    }

    public static void BuildAndroid()
    {
        var options = new BuildPlayerOptions
        {
            scenes = GetEnabledScenes(),
            locationPathName = "Builds/Android/Game.apk",
            target = BuildTarget.Android
        };

        BuildReport report = BuildPipeline.BuildPlayer(options);

        if (report.summary.result != BuildResult.Succeeded)
        {
            throw new Exception("Android build failed");
        }
    }
}
""",
)

add_heading(doc, "PowerShell Build Script Template", 1)
add_code_block(
    doc,
    r"""
$unityPath = "C:\Program Files\Unity\Hub\Editor\6000.0.66f2\Editor\Unity.exe"
$projectPath = "C:\Path\To\Your\Project"
$logPath = "$projectPath\Logs\android-build.log"

New-Item -ItemType Directory -Force -Path "$projectPath\Builds\Android" | Out-Null
New-Item -ItemType Directory -Force -Path "$projectPath\Logs" | Out-Null

& $unityPath `
  -quit `
  -batchmode `
  -projectPath $projectPath `
  -executeMethod BuildScript.BuildAndroid `
  -logFile $logPath

if ($LASTEXITCODE -ne 0) {
    Write-Host "Android build failed. Check log: $logPath"
    exit $LASTEXITCODE
}

Write-Host "Android build completed successfully."
Write-Host "APK saved in $projectPath\Builds\Android"
""",
)

add_heading(doc, "Button Launcher Template", 1)
add_code_block(
    doc,
    """
@echo off
powershell -ExecutionPolicy Bypass -File "%~dp0Build-Android.ps1"
pause
""",
)

add_heading(doc, "Optional Git Hook", 1)
doc.add_paragraph(
    "If you want automation on local git actions, use a local pre-push hook. "
    "For most teams, the one-click button is simpler and easier to control."
)
add_code_block(
    doc,
    """
#!/bin/sh
powershell -ExecutionPolicy Bypass -File "./BuildTools/Build-Android.ps1"
if [ $? -ne 0 ]; then
  echo "Build failed. Push cancelled."
  exit 1
fi
""",
)

add_heading(doc, "Outputs", 1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
header = table.rows[0].cells
header[0].text = "File"
header[1].text = "Suggested Path"
for cell in header:
    set_cell_shading(cell, "D9E7FF")

for name, path in [
    ("APK", r"Builds\Android\Game.apk"),
    ("Log", r"Logs\android-build.log"),
]:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = path

add_heading(doc, "Validation Checklist", 1)
for step in [
    "Double-click the local .bat file.",
    "Confirm Unity starts in batch mode.",
    "Confirm the APK is created in Builds/Android/.",
    "Confirm the log file is created in Logs/android-build.log.",
    "Install the APK and verify it runs.",
]:
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)

add_heading(doc, "Final Recommendation", 1)
doc.add_paragraph(
    "The best setup for your current goal is a local-only, Android-only, button-triggered build "
    "that saves the APK on your machine, with git-hook automation added only if you later decide you need it."
)

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
